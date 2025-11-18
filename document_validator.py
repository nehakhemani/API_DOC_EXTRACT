"""
Document Validator and Information Extractor

Analyzes PDF and Word documents to extract key information such as:
- Signature status
- Signing date
- Customer/client name
- Agreement type
- Pricing information

Supports: PDF, DOCX, DOC (with OCR for scanned PDFs)
"""

import os
import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
import logging

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentValidator:
    """Validates and extracts information from PDF and Word documents."""

    # Common signature indicators
    SIGNATURE_KEYWORDS = [
        r'signature',
        r'signed\s+by',
        r'electronically\s+signed',
        r'digitally\s+signed',
        r'/s/',
        r'authorized\s+signatory',
        r'executed\s+by',
        r'signed\s+on',
        r'date\s+signed',
        r'signature\s+date'
    ]

    # Signature role indicators (person with title/role and date suggests signature)
    SIGNATURE_ROLE_PATTERNS = [
        r'(?:client\s+lead|director|manager|president|ceo|cfo|authorized|signatory)',
        r'(?:vice\s+president|vp|secretary|treasurer|partner)',
        r'(?:representative|agent|officer|executive)',
    ]

    # Customer-side signature indicators
    CUSTOMER_ROLE_PATTERNS = [
        r'client\s+lead',
        r'customer\s+(?:representative|signatory|authorized)',
        r'buyer',
        r'purchaser',
        r'on\s+behalf\s+of\s+(?:client|customer)',
    ]

    # Spark NZ signature indicators
    SPARK_NZ_PATTERNS = [
        r'spark\s+(?:nz|new\s+zealand)',
        r'on\s+behalf\s+of\s+spark',
        r'for\s+and\s+on\s+behalf\s+of\s+spark',
        r'spark.*(?:representative|signatory|authorized)',
        r'account\s+(?:manager|executive)',
        r'sales\s+(?:manager|director|representative)',
    ]

    # Date patterns (various formats)
    DATE_PATTERNS = [
        r'\b(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\b',  # MM/DD/YYYY or DD/MM/YYYY
        r'\b(\d{4}[-/]\d{1,2}[-/]\d{1,2})\b',     # YYYY/MM/DD
        r'\b((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})\b',  # Month DD, YYYY
        r'\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})\b',  # DD Month YYYY
    ]

    # Agreement type patterns (order matters - more specific patterns first)
    AGREEMENT_TYPES = {
        'Variation': [r'agreement\s+variation', r'variation\s+to.*agreement', r'variation\s+number', r'variation\s+n[o°]?\.?:', r'variation\s+agreement'],
        'Statement of Work': [r'statement\s+of\s+work', r'SOW\b', r'work\s+statement'],
        'Letter of Engagement': [r'letter\s+of\s+engagement', r'engagement\s+letter'],
        'Service Schedule': [r'service\s+schedule', r'schedule\s+of\s+services'],
        'General Business Agreement': [r'general\s+business\s+agreement', r'business\s+agreement'],
        'Service Agreement': [r'service\s+agreement', r'service\s+contract'],
        'License Agreement': [r'license\s+agreement', r'licensing\s+agreement'],
        'NDA': [r'non[-\s]?disclosure\s+agreement', r'confidentiality\s+agreement'],
        'Sales Agreement': [r'sales\s+agreement', r'purchase\s+agreement'],
        'Master Agreement': [r'master\s+agreement', r'master\s+service\s+agreement', r'MSA'],
    }

    def __init__(self, log_level: str = "INFO", use_ocr: bool = False):
        """
        Initialize the document validator.

        Args:
            log_level: Logging level (DEBUG, INFO, WARNING, ERROR)
            use_ocr: Enable OCR for scanned PDFs (requires Tesseract/Poppler installation)
        """
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(getattr(logging, log_level))
        self.use_ocr = use_ocr

        if not self.logger.handlers:
            handler = logging.StreamHandler()
            handler.setFormatter(logging.Formatter('%(levelname)s - %(message)s'))
            self.logger.addHandler(handler)

        self._check_dependencies()

    def _check_dependencies(self):
        """Check which document libraries are available."""
        if not PYPDF2_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            self.logger.warning("No PDF libraries available. Install PyPDF2 or pdfplumber for PDF support.")
        elif PDFPLUMBER_AVAILABLE:
            self.logger.info("Using pdfplumber for PDF extraction")
        elif PYPDF2_AVAILABLE:
            self.logger.info("Using PyPDF2 for PDF extraction")

        if self.use_ocr:
            if PDF2IMAGE_AVAILABLE and PYTESSERACT_AVAILABLE:
                self.logger.info("OCR enabled and available for scanned PDFs")
            else:
                self.logger.warning("OCR enabled but libraries not available!")
                self.logger.warning("Install: pip install pdf2image pytesseract")
                self.logger.warning("And install Tesseract-OCR and Poppler (see OCR_SETUP.md)")
        else:
            self.logger.info("OCR disabled - scanned PDFs will fail with error message")

        if DOCX_AVAILABLE:
            self.logger.info("Word document support available (.docx)")
        else:
            self.logger.warning("Word document support not available. Install python-docx for .docx support.")

    def extract_text_pypdf2(self, pdf_path: str) -> str:
        """Extract text from PDF using PyPDF2."""
        if not PYPDF2_AVAILABLE:
            return ""

        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            self.logger.error(f"PyPDF2 extraction failed: {e}")

        return text

    def extract_text_pdfplumber(self, pdf_path: str) -> str:
        """Extract text from PDF using pdfplumber."""
        if not PDFPLUMBER_AVAILABLE:
            return ""

        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            self.logger.error(f"pdfplumber extraction failed: {e}")

        return text

    def extract_text_ocr(self, pdf_path: str, max_pages: int = 10) -> str:
        """
        Extract text from image-based/scanned PDF using OCR.

        Args:
            pdf_path: Path to PDF file
            max_pages: Maximum number of pages to OCR (default: 10 for better coverage)

        Returns:
            Extracted text
        """
        if not PDF2IMAGE_AVAILABLE or not PYTESSERACT_AVAILABLE:
            self.logger.warning("OCR libraries not available.")
            if not PDF2IMAGE_AVAILABLE:
                self.logger.warning("→ Missing: pdf2image (pip install pdf2image)")
            if not PYTESSERACT_AVAILABLE:
                self.logger.warning("→ Missing: pytesseract (pip install pytesseract)")
            self.logger.warning("→ See OCR_SETUP.md for complete setup instructions")
            return ""

        text = ""
        try:
            self.logger.info(f"Performing OCR on scanned PDF (processing up to {max_pages} pages)...")

            # Convert PDF to images with higher DPI for better quality
            images = convert_from_path(pdf_path, dpi=300, first_page=1, last_page=max_pages)

            # Perform OCR on each page with enhanced configuration
            for i, image in enumerate(images):
                self.logger.info(f"OCR processing page {i+1}/{len(images)}...")

                # Preprocess image for better OCR accuracy
                # Convert to grayscale and enhance contrast
                from PIL import ImageEnhance
                # Convert to grayscale
                image = image.convert('L')
                # Enhance contrast
                enhancer = ImageEnhance.Contrast(image)
                image = enhancer.enhance(2.0)
                # Enhance sharpness
                enhancer = ImageEnhance.Sharpness(image)
                image = enhancer.enhance(1.5)

                # OCR with custom config for better accuracy
                custom_config = r'--oem 3 --psm 6'
                page_text = pytesseract.image_to_string(image, config=custom_config)
                if page_text:
                    text += page_text + "\n"

            self.logger.info(f"OCR complete. Extracted {len(text)} characters from {len(images)} pages.")

        except Exception as e:
            error_msg = str(e)
            self.logger.error(f"OCR extraction failed: {error_msg}")

            # Provide helpful guidance based on error type
            if "poppler" in error_msg.lower() or "Unable to get page count" in error_msg:
                self.logger.error("→ Poppler is not installed or not in PATH")
                self.logger.error("→ See OCR_SETUP.md for installation instructions")
            elif "tesseract" in error_msg.lower():
                self.logger.error("→ Tesseract OCR is not installed or not in PATH")
                self.logger.error("→ See OCR_SETUP.md for installation instructions")
            else:
                self.logger.error("→ Check OCR_SETUP.md for troubleshooting")

        return text

    def extract_text(self, pdf_path: str, min_text_threshold: int = 100) -> str:
        """
        Extract text from PDF using available libraries.
        Tries pdfplumber first, falls back to PyPDF2, then OCR if enabled.

        Args:
            pdf_path: Path to PDF file
            min_text_threshold: Minimum characters before considering extraction successful (default: 100)

        Returns:
            Extracted text, or empty string if extraction failed
        """
        if not os.path.exists(pdf_path):
            self.logger.error(f"File not found: {pdf_path}")
            return ""

        # Try pdfplumber first (generally better)
        if PDFPLUMBER_AVAILABLE:
            text = self.extract_text_pdfplumber(pdf_path)
            if len(text.strip()) >= min_text_threshold:
                return text

        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            text = self.extract_text_pypdf2(pdf_path)
            if len(text.strip()) >= min_text_threshold:
                return text

        # If we have some text but not enough, it's likely a scanned/image-based PDF
        # Check if OCR is enabled
        if self.use_ocr:
            if PDF2IMAGE_AVAILABLE and PYTESSERACT_AVAILABLE:
                self.logger.info("Insufficient text extracted, trying OCR...")
                ocr_text = self.extract_text_ocr(pdf_path)
                if ocr_text.strip():
                    return ocr_text
            else:
                self.logger.error("OCR enabled but libraries not installed!")
                self.logger.error("→ Install: pip install pdf2image pytesseract")
                self.logger.error("→ System requirements: Tesseract-OCR and Poppler")
                self.logger.error("→ See OCR_SETUP.md for installation instructions")
                return ""
        else:
            # OCR is disabled but document appears to be scanned
            extracted_chars = len(text.strip()) if 'text' in locals() else 0
            self.logger.error(f"Document appears to be scanned/image-based (only {extracted_chars} characters extracted)")
            self.logger.error("→ This document requires OCR processing")
            self.logger.error("→ Enable OCR in batch_config.json: \"use_ocr\": true")
            self.logger.error("→ Alternative: Use cloud-based OCR services (Azure AI Document Intelligence, AWS Textract, Google Cloud Vision)")
            return ""

        return ""

    def extract_text_docx(self, docx_path: str) -> str:
        """
        Extract text from DOCX file.

        Args:
            docx_path: Path to DOCX file

        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx not available. Install with: pip install python-docx")
            return ""

        text = ""
        try:
            doc = Document(docx_path)

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            self.logger.info(f"Extracted {len(text)} characters from DOCX")

        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")

        return text

    def extract_text_from_document(self, file_path: str) -> str:
        """
        Extract text from any supported document format (PDF, DOCX, DOC).

        Args:
            file_path: Path to document file

        Returns:
            Extracted text
        """
        if not os.path.exists(file_path):
            self.logger.error(f"File not found: {file_path}")
            return ""

        file_ext = Path(file_path).suffix.lower()

        # Handle Word documents
        if file_ext in ['.docx', '.doc']:
            if file_ext == '.docx':
                return self.extract_text_docx(file_path)
            elif file_ext == '.doc':
                self.logger.warning(".doc files (old Word format) are not directly supported.")
                self.logger.warning("Please convert to .docx or save as .docx format.")
                return ""

        # Handle PDF documents
        elif file_ext == '.pdf':
            return self.extract_text(file_path)

        else:
            self.logger.error(f"Unsupported file format: {file_ext}")
            self.logger.error("Supported formats: .pdf, .docx")
            return ""

    def detect_signature(self, text: str) -> Dict[str, Any]:
        """
        Detect if document appears to be signed based on keywords and patterns.

        Returns:
            dict with 'is_signed' (bool), 'confidence' (str), and 'indicators' (list)
        """
        text_lower = text.lower()
        found_indicators = []

        # Check for explicit signature keywords
        for pattern in self.SIGNATURE_KEYWORDS:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                found_indicators.append(match.group())

        # Check for signature roles (Client Lead, Director, etc.) near dates
        # This indicates a signature block even without the word "signature"
        role_found = False
        for pattern in self.SIGNATURE_ROLE_PATTERNS:
            if re.search(pattern, text_lower, re.IGNORECASE):
                role_found = True
                found_indicators.append(f"role: {pattern.split('|')[0].replace('(?:', '').strip()}")
                break

        # If we find a role/title AND a date, it's likely a signature block
        dates = self.extract_dates(text)
        has_date = len(dates) > 0

        # Determine confidence level
        indicator_count = len(found_indicators)

        # Role + Date is a strong indicator of signature even without "signature" keyword
        if role_found and has_date:
            if indicator_count >= 1:
                confidence = "high"
                is_signed = True
            else:
                confidence = "medium"
                is_signed = True
                found_indicators.append("role with date (signature block detected)")
        elif indicator_count >= 3:
            confidence = "high"
            is_signed = True
        elif indicator_count >= 1:
            confidence = "medium"
            is_signed = True
        else:
            confidence = "low"
            is_signed = False

        return {
            'is_signed': is_signed,
            'confidence': confidence,
            'indicator_count': indicator_count,
            'indicators': list(set(found_indicators))[:5]  # Unique, max 5
        }

    def detect_signatories(self, text: str) -> Dict[str, Any]:
        """
        Detect and classify signatories as customer or Spark NZ.

        Returns:
            dict with customer and spark_nz signature information
        """
        text_lower = text.lower()
        lines = text.split('\n')

        customer_signed = False
        spark_nz_signed = False
        customer_info = None
        spark_nz_info = None

        # Check for customer signature
        for pattern in self.CUSTOMER_ROLE_PATTERNS:
            match = re.search(pattern, text_lower, re.IGNORECASE)
            if match:
                customer_signed = True
                # Try to extract name and date near this role
                for i, line in enumerate(lines):
                    if re.search(pattern, line.lower(), re.IGNORECASE):
                        # Look for name in nearby lines (usually name comes before role)
                        name = None
                        date = None

                        # Check previous line for name
                        if i > 0:
                            potential_name = lines[i-1].strip()
                            # Name should be 2-50 chars, letters and spaces
                            if len(potential_name) > 2 and len(potential_name) < 50 and potential_name[0].isupper():
                                name = potential_name

                        # Check next few lines for date
                        for j in range(i, min(i+3, len(lines))):
                            line_dates = self.extract_dates(lines[j])
                            if line_dates:
                                date = line_dates[0]
                                break

                        customer_info = {
                            'signed': True,
                            'name': name,
                            'role': match.group(),
                            'date': date
                        }
                        break
                break

        # Check for Spark NZ signature
        for pattern in self.SPARK_NZ_PATTERNS:
            match = re.search(pattern, text_lower, re.IGNORECASE)
            if match:
                spark_nz_signed = True
                # Try to extract name and date near this role
                for i, line in enumerate(lines):
                    if re.search(pattern, line.lower(), re.IGNORECASE):
                        # Look for name in nearby lines
                        name = None
                        date = None

                        if i > 0:
                            potential_name = lines[i-1].strip()
                            if len(potential_name) > 2 and len(potential_name) < 50 and potential_name[0].isupper():
                                name = potential_name

                        for j in range(i, min(i+3, len(lines))):
                            line_dates = self.extract_dates(lines[j])
                            if line_dates:
                                date = line_dates[0]
                                break

                        spark_nz_info = {
                            'signed': True,
                            'name': name,
                            'role': match.group(),
                            'date': date
                        }
                        break
                break

        return {
            'customer': customer_info or {'signed': False, 'name': None, 'role': None, 'date': None},
            'spark_nz': spark_nz_info or {'signed': False, 'name': None, 'role': None, 'date': None},
            'both_signed': customer_signed and spark_nz_signed
        }

    def extract_dates(self, text: str) -> List[str]:
        """Extract potential dates from document text."""
        dates = []

        for pattern in self.DATE_PATTERNS:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                dates.append(match.group(1))

        # Remove duplicates while preserving order
        seen = set()
        unique_dates = []
        for date in dates:
            if date not in seen:
                seen.add(date)
                unique_dates.append(date)

        return unique_dates

    def extract_signing_date(self, text: str) -> Optional[str]:
        """
        Extract the most likely signing date from document.
        Looks for dates near signature-related keywords.
        """
        lines = text.split('\n')
        signing_date = None

        # Look for dates near signature keywords
        for i, line in enumerate(lines):
            line_lower = line.lower()

            # Check if line contains signature-related keywords
            for keyword_pattern in ['date.*signed', 'signed.*date', 'signature.*date', 'executed.*date']:
                if re.search(keyword_pattern, line_lower):
                    # Look for date in this line and next few lines
                    context = '\n'.join(lines[i:min(i+3, len(lines))])
                    dates = self.extract_dates(context)
                    if dates:
                        signing_date = dates[0]
                        break

            if signing_date:
                break

        # If no signing date found, return the last date in document (common pattern)
        if not signing_date:
            all_dates = self.extract_dates(text)
            if all_dates:
                signing_date = all_dates[-1]

        return signing_date

    def extract_customer_name(self, text: str, filename: str = "") -> Optional[str]:
        """
        Extract customer/client name from document.
        Uses filename parsing and text patterns.
        """
        # Try to extract from filename first (common pattern: ID_CUSTOMERNAME_DocumentType.pdf)
        if filename:
            parts = filename.replace('.pdf', '').replace('.PDF', '').split('_')
            if len(parts) >= 2:
                # Second part is usually customer name
                customer_name = parts[1].strip()
                if customer_name and len(customer_name) > 2:
                    return customer_name.replace('_', ' ').title()

        # Try to find in document text
        patterns = [
            r'(?:client|customer|company)\s*(?:name)?:\s*([A-Z][A-Za-z\s&.,]+(?:Ltd|LLC|Inc|Corp|Limited|Co)?)',
            r'(?:between|with)\s+([A-Z][A-Za-z\s&.,]+(?:Ltd|LLC|Inc|Corp|Limited|Co)?)\s+(?:and|&)',
            r'This\s+agreement.*?(?:between|with)\s+([A-Z][A-Za-z\s&.,]+)',
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                # Clean up and validate
                if len(name) > 3 and len(name) < 100:
                    return name

        return None

    def detect_agreement_type(self, text: str, filename: str = "") -> Dict[str, Any]:
        """
        Detect the type of agreement from document text and filename.

        Returns:
            dict with 'type', 'confidence', and 'matched_pattern'
        """
        text_lower = text.lower()
        combined_text = f"{filename} {text_lower}"

        matches = []

        for agreement_type, patterns in self.AGREEMENT_TYPES.items():
            for pattern in patterns:
                if re.search(pattern, combined_text, re.IGNORECASE):
                    # Count occurrences for confidence
                    count = len(re.findall(pattern, combined_text, re.IGNORECASE))
                    matches.append({
                        'type': agreement_type.replace('_', ' ').title(),
                        'pattern': pattern,
                        'count': count
                    })

        if not matches:
            return {'type': 'Unknown', 'confidence': 'low', 'matched_pattern': None}

        # Sort by count and pick the best match
        best_match = sorted(matches, key=lambda x: x['count'], reverse=True)[0]

        confidence = 'high' if best_match['count'] >= 2 else 'medium'

        return {
            'type': best_match['type'],
            'confidence': confidence,
            'matched_pattern': best_match['pattern']
        }

    def extract_pricing(self, text: str) -> Dict[str, Any]:
        """
        Extract pricing information from document text.

        Returns:
            dict with pricing details found in the document
        """
        pricing_info = {
            'has_pricing': False,
            'amounts': [],
            'pricing_section_found': False
        }

        # Check if document has a pricing section
        if re.search(r'^\s*\d+\.\s*(?:pricing|fees|charges|cost)', text, re.MULTILINE | re.IGNORECASE):
            pricing_info['pricing_section_found'] = True

        # Extract all monetary amounts
        price_patterns = [
            r'\$\s*[\d,]+\.?\d*\s*(?:per|/)\s*(?:connection|month|user|line|year|week|day)',
            r'\$\s*[\d,]+\.?\d*',
            r'(?:AUD|NZD|USD|EUR|GBP)\s*[\d,]+\.?\d*',
        ]

        amounts = []
        for pattern in price_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                amount = match.group().strip()
                if amount and amount not in amounts:
                    amounts.append(amount)

        if amounts:
            pricing_info['has_pricing'] = True
            pricing_info['amounts'] = amounts[:10]  # Limit to first 10

        return pricing_info

    def validate_document(self, file_path: str) -> Dict[str, Any]:
        """
        Perform complete validation and information extraction on a document.

        Args:
            file_path: Path to the document file (PDF, DOCX)

        Returns:
            Dictionary containing all extracted information
        """
        self.logger.info(f"Validating document: {file_path}")

        # Get filename and file type
        filename = os.path.basename(file_path)
        file_ext = Path(file_path).suffix.lower()

        # Extract text based on file type
        text = self.extract_text_from_document(file_path)

        if not text.strip():
            return {
                'filename': filename,
                'status': 'error',
                'error': f'Could not extract text from {file_ext} file',
                'file_path': file_path
            }

        # Perform all extractions
        signature_info = self.detect_signature(text)
        signatories_info = self.detect_signatories(text)
        signing_date = self.extract_signing_date(text)
        customer_name = self.extract_customer_name(text, filename)
        agreement_info = self.detect_agreement_type(text, filename)
        all_dates = self.extract_dates(text)
        pricing_info = self.extract_pricing(text)

        # Compile results
        results = {
            'filename': filename,
            'file_path': file_path,
            'status': 'success',
            'signature': {
                'is_signed': signature_info['is_signed'],
                'confidence': signature_info['confidence'],
                'indicators_found': signature_info['indicator_count'],
                'signature_indicators': signature_info['indicators']
            },
            'signatories': {
                'customer': signatories_info['customer'],
                'spark_nz': signatories_info['spark_nz'],
                'both_signed': signatories_info['both_signed']
            },
            'signing_date': signing_date,
            'customer_name': customer_name,
            'agreement_type': {
                'type': agreement_info['type'],
                'confidence': agreement_info['confidence']
            },
            'pricing': {
                'has_pricing': pricing_info['has_pricing'],
                'pricing_section_found': pricing_info['pricing_section_found'],
                'amounts': pricing_info['amounts']
            },
            'extracted_dates': all_dates[:10],  # Limit to first 10 dates
            'text_length': len(text),
            'analyzed_at': datetime.now().isoformat()
        }

        self.logger.info(f"Validation complete: {results['agreement_type']['type']} - Signed: {results['signature']['is_signed']}")

        return results

    def validate_directory(self, directory: str, output_file: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        Validate all PDF files in a directory.

        Args:
            directory: Path to directory containing PDFs
            output_file: Optional JSON file path to save results

        Returns:
            List of validation results for each PDF
        """
        results = []
        pdf_files = list(Path(directory).glob('*.pdf')) + list(Path(directory).glob('*.PDF'))

        self.logger.info(f"Found {len(pdf_files)} PDF files in {directory}")

        for pdf_path in pdf_files:
            result = self.validate_document(str(pdf_path))
            results.append(result)

        # Save to JSON if requested
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            self.logger.info(f"Results saved to {output_file}")

        return results

    def print_summary(self, result: Dict[str, Any]):
        """Print a formatted summary of validation results."""
        print("\n" + "="*60)
        print(f"Document: {result['filename']}")
        print("="*60)

        if result['status'] == 'error':
            print(f"Error: {result['error']}")
            return

        print(f"\nAgreement Type: {result['agreement_type']['type']} ({result['agreement_type']['confidence']} confidence)")
        print(f"Customer Name: {result['customer_name'] or 'Not detected'}")
        print(f"Signed: {'Yes' if result['signature']['is_signed'] else 'No'} ({result['signature']['confidence']} confidence)")
        print(f"Signing Date: {result['signing_date'] or 'Not detected'}")

        # Display detailed signatory information
        if 'signatories' in result:
            print(f"\n--- SIGNATORIES ---")

            customer_sig = result['signatories']['customer']
            if customer_sig['signed']:
                print(f"Customer Signature: YES")
                if customer_sig['name']:
                    print(f"  Name: {customer_sig['name']}")
                if customer_sig['role']:
                    print(f"  Role: {customer_sig['role']}")
                if customer_sig['date']:
                    print(f"  Date: {customer_sig['date']}")
            else:
                print(f"Customer Signature: NO")

            spark_nz_sig = result['signatories']['spark_nz']
            if spark_nz_sig['signed']:
                print(f"Spark NZ Signature: YES")
                if spark_nz_sig['name']:
                    print(f"  Name: {spark_nz_sig['name']}")
                if spark_nz_sig['role']:
                    print(f"  Role: {spark_nz_sig['role']}")
                if spark_nz_sig['date']:
                    print(f"  Date: {spark_nz_sig['date']}")
            else:
                print(f"Spark NZ Signature: NO")

            if result['signatories']['both_signed']:
                print(f"\nBoth parties have signed: YES")

        if result['signature']['signature_indicators']:
            print(f"\nSignature Indicators Found: {', '.join(result['signature']['signature_indicators'][:3])}")

        if result['extracted_dates']:
            print(f"\nDates Found: {', '.join(result['extracted_dates'][:5])}")

        print("\n" + "="*60)


def main():
    """Command-line interface for document validation."""
    import argparse

    parser = argparse.ArgumentParser(description='Validate and extract information from PDF documents')
    parser.add_argument('path', help='PDF file or directory path')
    parser.add_argument('--output', '-o', help='Output JSON file for results')
    parser.add_argument('--log-level', default='INFO', choices=['DEBUG', 'INFO', 'WARNING', 'ERROR'],
                       help='Logging level')

    args = parser.parse_args()

    validator = DocumentValidator(log_level=args.log_level)

    if os.path.isfile(args.path):
        # Single file
        result = validator.validate_document(args.path)
        validator.print_summary(result)

        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            print(f"\nResults saved to {args.output}")

    elif os.path.isdir(args.path):
        # Directory
        results = validator.validate_directory(args.path, args.output)

        print(f"\n{'='*60}")
        print(f"Processed {len(results)} documents")
        print(f"{'='*60}")

        # Summary statistics
        signed_count = sum(1 for r in results if r.get('signature', {}).get('is_signed'))
        print(f"\nSigned documents: {signed_count}/{len(results)}")

        # Agreement types
        types = {}
        for r in results:
            atype = r.get('agreement_type', {}).get('type', 'Unknown')
            types[atype] = types.get(atype, 0) + 1

        print("\nAgreement Types:")
        for atype, count in sorted(types.items(), key=lambda x: x[1], reverse=True):
            print(f"  {atype}: {count}")

    else:
        print(f"Error: {args.path} is not a valid file or directory")


if __name__ == '__main__':
    main()
